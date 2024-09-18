import streamlit as st
from docx import Document
import io
import re

# Set page configuration
st.set_page_config(page_title="Text to Word Converter", layout="wide")

# Logo
st.image("jb.png", width=250)  # Replace with your logo URL

# Navbar
st.markdown("""
<div class="navbar">
  <a href="#Home">Home</a>
  <a href="#About">About</a>
  <a href="https://techiehelpt.netlify.app/">Back To Website</a>
</div>
""", unsafe_allow_html=True)

# Add navigation to different sections
menu = ["Home", "About"]
choice = st.sidebar.selectbox("Navigate", menu)

if choice == "Home":
    # Home Page
    st.title("Text to Word Converter")
    st.write("Enter your text below to convert it into a Word document.")

    # Text input
    text_input = st.text_area("Enter text here", height=300)

    if st.button("Convert to Word"):
        if text_input:
            # Create a Word document
            doc = Document()
            doc.add_heading('Converted Text', level=1)
            doc.add_paragraph(text_input)

            # Save the document to a BytesIO stream
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            # Download button for the Word document
            st.download_button(
                label="📥 Download Word Document",
                data=buffer,
                file_name="converted_text.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        else:
            st.error("Please enter some text before converting to Word.")

elif choice == "About":
    # About Page
    st.title("About Text to Word Converter")
    st.write("""
    This tool allows you to convert plain text into a Word document.

    **Features**:
    - Enter text in a text area.
    - Convert the text into a Word document.
    - Download the generated Word document.

    **Technology Stack**:
    - Streamlit (for building the web app)
    - python-docx (for creating Word documents)
    """)

# Footer
st.markdown("""
<div class="footer">
    <p>© 2024 Text to Word Converter | TechieHelp</p>
    <a href="https://www.linkedin.com/in/techiehelp" style="color:white; margin-right: 10px;">LinkedIn</a>
    <a href="https://www.twitter.com/techiehelp" style="color:white; margin-right: 10px;">Twitter</a>
    <a href="https://www.instagram.com/techiehelp2" style="color:white;">Instagram</a>
</div>
""", unsafe_allow_html=True)
